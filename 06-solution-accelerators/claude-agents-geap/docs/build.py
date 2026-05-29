"""Full engineering design document builder.

Output: eddoc/cc-claude-code-on-ge-engineering-reference.docx
Audience: senior engineers, solution architects, partner engineering leads.

Style baseline extracted from
~/Documents/claude-code-vertex-gcp-engineering-design.docx
(palette + Calibri + colored-table layout idiom).

Layout: two sections.
  Section 1 (cover): zero margins, no header/footer.
  Section 2 (body): standard margins, running header + footer with page
                    numbers, restarting at 1.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL

# ============================================================================
# Palette (verified vs reference doc)
# ============================================================================
COPPER       = "D97757"
COPPER_DARK  = "CC785C"
COPPER_BG    = "FBE9DF"
WARM_WHITE   = "FAF9F5"
WARM_LIGHT   = "F8F4ED"
BLUE         = "1A73E8"
BLUE_BG      = "E8F0FE"
GREEN        = "34A853"
GREEN_BG     = "E6F4EA"
GRAY_MUTED   = "5F6368"
NEAR_BLACK   = "141413"
BODY_GRAY    = "202124"
LIGHT_GRAY   = "F1F3F4"
WHITE        = "FFFFFF"

FONT      = "Calibri"
FONT_MONO = "Consolas"


# ============================================================================
# Low-level OOXML helpers
# ============================================================================

def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_width(cell, width_in):
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_in * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcPr.append(tcW)


def set_row_height(row, height_in, rule='atLeast'):
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_in * 1440)))
    trHeight.set(qn('w:hRule'), rule)
    trPr = row._tr.get_or_add_trPr()
    for existing in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing)
    trPr.append(trHeight)


def set_cell_margins(cell, top_in=0, bottom_in=0, left_in=0, right_in=0):
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(existing)
    tcMar = OxmlElement('w:tcMar')
    for direction, val in [('top', top_in), ('left', left_in),
                            ('bottom', bottom_in), ('right', right_in)]:
        m = OxmlElement(f'w:{direction}')
        m.set(qn('w:w'), str(int(val * 1440)))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def vert_align_cell(cell, val='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(existing)
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), val)
    tcPr.append(vAlign)


def remove_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def remove_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_thin_table_borders(tbl, hex_color=LIGHT_GRAY):
    tblPr = tbl._tbl.tblPr
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), hex_color)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    if line_pt is not None:
        pf.line_spacing = Pt(line_pt)


def add_run(p, text, *, color=BODY_GRAY, size=11, bold=False, italic=False,
            font=FONT, caps=False, char_spacing=None):
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    if italic:
        run.italic = True
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('ascii', 'hAnsi', 'cs', 'eastAsia'):
        rFonts.set(qn(f'w:{attr}'), font)
    if caps:
        caps_el = OxmlElement('w:caps')
        rPr.append(caps_el)
    if char_spacing is not None:
        spacing_el = OxmlElement('w:spacing')
        spacing_el.set(qn('w:val'), str(int(char_spacing * 20)))
        rPr.append(spacing_el)
    return run


def add_bottom_border(p, hex_color, size_eighths=6):
    pPr = p._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:pBdr')):
        pPr.remove(existing)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size_eighths))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_page_break_paragraph(container):
    """For containers that aren't the doc itself (cells, etc.)."""
    p = container.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


# ============================================================================
# High-level content builders
# ============================================================================

def add_h1(doc, number, text):
    """Section heading: blue, 20pt bold."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=24, after_pt=12, line_pt=26)
    add_run(p, f"{number}.  {text}", color=BLUE, size=20, bold=True)
    return p


def add_h2(doc, number, text):
    """Subsection heading: near-black, 14pt bold."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=14, after_pt=6, line_pt=18)
    add_run(p, f"{number}  {text}", color=NEAR_BLACK, size=14, bold=True)
    return p


def add_h3(doc, text):
    """Sub-sub heading: near-black, 12pt bold."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=10, after_pt=4, line_pt=16)
    add_run(p, text, color=NEAR_BLACK, size=12, bold=True)
    return p


def add_para(doc, text, *, size=11, color=BODY_GRAY, italic=False):
    """Body paragraph."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=8, line_pt=16)
    add_run(p, text, color=color, size=size, italic=italic)
    return p


def add_para_runs(doc, parts):
    """Body paragraph with mixed-formatting runs.
    parts: list of (text, options-dict-or-None)
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=8, line_pt=16)
    for item in parts:
        if isinstance(item, str):
            add_run(p, item, color=BODY_GRAY, size=11)
        else:
            text, opts = item
            kwargs = dict(color=BODY_GRAY, size=11)
            kwargs.update(opts or {})
            add_run(p, text, **kwargs)
    return p


def add_bullet(doc, text, *, size=11, color=BODY_GRAY):
    """List bullet (uses Word's List Bullet style if present, else manual)."""
    p = doc.add_paragraph(style='List Bullet')
    set_paragraph_spacing(p, before_pt=0, after_pt=4, line_pt=16)
    add_run(p, text, color=color, size=size)
    return p


def add_callout(doc, label, body, *, label_color=COPPER, bg=COPPER_BG):
    """Copper-bordered callout (1x1 table with tinted bg + left accent)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.allow_autofit = False
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9360')  # 6.5"
    tblW.set(qn('w:type'), 'dxa')
    tbl._tbl.tblPr.append(tblW)
    remove_table_borders(tbl)

    cell = tbl.rows[0].cells[0]
    set_cell_width(cell, 6.5)
    set_cell_shading(cell, bg)
    set_cell_margins(cell, top_in=0.12, bottom_in=0.12,
                     left_in=0.2, right_in=0.2)
    remove_cell_borders(cell)

    # Apply a left border accent (colored bar) via cell borders
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')  # 3pt
    left.set(qn('w:color'), label_color)
    tcBorders.append(left)
    for side in ('top', 'right', 'bottom'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=16)
    add_run(p, f"{label}  ", color=label_color, size=10, bold=True,
            char_spacing=1.2)
    add_run(p, body, color=BODY_GRAY, size=10.5)

    # Spacer paragraph after callout
    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before_pt=4, after_pt=4)


def add_code_block(doc, code, *, mono_size=9.5):
    """Monospace code block in a tinted-background cell."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.allow_autofit = False
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9360')
    tblW.set(qn('w:type'), 'dxa')
    tbl._tbl.tblPr.append(tblW)
    remove_table_borders(tbl)

    cell = tbl.rows[0].cells[0]
    set_cell_width(cell, 6.5)
    set_cell_shading(cell, WARM_LIGHT)
    set_cell_margins(cell, top_in=0.1, bottom_in=0.1,
                     left_in=0.15, right_in=0.15)
    remove_cell_borders(cell)

    lines = code.split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=13)
        add_run(p, line or " ", color=NEAR_BLACK, size=mono_size,
                font=FONT_MONO)

    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before_pt=4, after_pt=4)


def add_table_styled(doc, headers, rows, *, col_widths_in=None):
    """Standard table: blue header row, alternating warm-light body rows."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.autofit = False
    tbl.allow_autofit = False
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9360')
    tblW.set(qn('w:type'), 'dxa')
    tbl._tbl.tblPr.append(tblW)
    add_thin_table_borders(tbl, LIGHT_GRAY)

    # Header row
    hdr = tbl.rows[0]
    for ci, h in enumerate(headers):
        cell = hdr.cells[ci]
        if col_widths_in:
            set_cell_width(cell, col_widths_in[ci])
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell, top_in=0.06, bottom_in=0.06,
                         left_in=0.1, right_in=0.1)
        vert_align_cell(cell, 'center')
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=14)
        add_run(p, h, color=WHITE, size=10.5, bold=True, char_spacing=0.5)

    # Body rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[1 + ri]
        bg = WARM_LIGHT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            if col_widths_in:
                set_cell_width(cell, col_widths_in[ci])
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top_in=0.06, bottom_in=0.06,
                             left_in=0.1, right_in=0.1)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=14)
            add_run(p, val, color=BODY_GRAY, size=10)

    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before_pt=4, after_pt=4)


# ============================================================================
# Cover (carried over from build_cover.py)
# ============================================================================

def build_cover(doc):
    cover = doc.add_table(rows=1, cols=4)
    cover.autofit = False
    cover.allow_autofit = False
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '12240')
    tblW.set(qn('w:type'), 'dxa')
    cover._tbl.tblPr.append(tblW)
    remove_table_borders(cover)

    row = cover.rows[0]
    set_row_height(row, 10.7, rule='atLeast')
    cells = row.cells
    copper_bar, blue_bar, gutter, content = cells

    set_cell_width(copper_bar, 0.18)
    set_cell_shading(copper_bar, COPPER)
    set_cell_margins(copper_bar, 0, 0, 0, 0)
    remove_cell_borders(copper_bar)

    set_cell_width(blue_bar, 0.18)
    set_cell_shading(blue_bar, BLUE)
    set_cell_margins(blue_bar, 0, 0, 0, 0)
    remove_cell_borders(blue_bar)

    set_cell_width(gutter, 0.3)
    set_cell_margins(gutter, 0, 0, 0, 0)
    remove_cell_borders(gutter)

    set_cell_width(content, 7.84)
    set_cell_margins(content, top_in=0.7, left_in=0.4,
                     right_in=0.5, bottom_in=0.5)
    remove_cell_borders(content)

    p0 = content.paragraphs[0]
    p0.text = ""
    set_paragraph_spacing(p0, 0, 0)

    pill_tbl = content.add_table(rows=1, cols=3)
    pill_tbl.autofit = False
    pill_tbl.allow_autofit = False
    pill_row = pill_tbl.rows[0]
    set_row_height(pill_row, 0.32, rule='atLeast')
    remove_table_borders(pill_tbl)
    p_anth, p_gap, p_gc = pill_row.cells

    set_cell_width(p_anth, 1.45)
    set_cell_shading(p_anth, COPPER)
    set_cell_margins(p_anth, top_in=0.04, bottom_in=0.04,
                     left_in=0.15, right_in=0.15)
    vert_align_cell(p_anth, 'center')
    remove_cell_borders(p_anth)
    pa = p_anth.paragraphs[0]
    pa.text = ""
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(pa, 0, 0)
    add_run(pa, "ANTHROPIC", color=WHITE, size=10, bold=True, char_spacing=1.6)

    set_cell_width(p_gap, 0.18)
    set_cell_margins(p_gap, 0, 0, 0, 0)
    remove_cell_borders(p_gap)

    set_cell_width(p_gc, 1.6)
    set_cell_shading(p_gc, BLUE)
    set_cell_margins(p_gc, top_in=0.04, bottom_in=0.04,
                     left_in=0.15, right_in=0.15)
    vert_align_cell(p_gc, 'center')
    remove_cell_borders(p_gc)
    pg = p_gc.paragraphs[0]
    pg.text = ""
    pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(pg, 0, 0)
    add_run(pg, "GOOGLE CLOUD", color=WHITE, size=10, bold=True,
            char_spacing=1.6)

    p_eye = content.add_paragraph()
    set_paragraph_spacing(p_eye, before_pt=40, after_pt=0)
    add_run(p_eye, "CO-INNOVATION PROPOSAL", color=BLUE, size=10, bold=True,
            char_spacing=2.2)

    p_t1 = content.add_paragraph()
    set_paragraph_spacing(p_t1, before_pt=14, after_pt=0, line_pt=52)
    add_run(p_t1, "Claude Powered Agents", color=NEAR_BLACK, size=44, bold=True)

    p_t2 = content.add_paragraph()
    set_paragraph_spacing(p_t2, before_pt=0, after_pt=0, line_pt=52)
    add_run(p_t2, "for Google Cloud", color=NEAR_BLACK, size=44, bold=True)

    p_sub = content.add_paragraph()
    set_paragraph_spacing(p_sub, before_pt=18, after_pt=0, line_pt=20)
    add_run(p_sub, "Claude-Powered AI Agents on the Gemini Enterprise Agent "
                   "Platform & Google Cloud Marketplace",
            color=GRAY_MUTED, size=14)

    p_rule = content.add_paragraph()
    set_paragraph_spacing(p_rule, before_pt=30, after_pt=0)
    add_bottom_border(p_rule, LIGHT_GRAY, size_eighths=8)

    p_pb = content.add_paragraph()
    set_paragraph_spacing(p_pb, before_pt=20, after_pt=0)
    add_run(p_pb, "Prepared by", color=GRAY_MUTED, size=10, char_spacing=0.5)

    p_name = content.add_paragraph()
    set_paragraph_spacing(p_name, before_pt=6, after_pt=0)
    add_run(p_name, "Schneider Larbi", color=NEAR_BLACK, size=20, bold=True)

    p_role = content.add_paragraph()
    set_paragraph_spacing(p_role, before_pt=6, after_pt=0, line_pt=16)
    add_run(p_role, "Senior Manager, Global Partner Technical Architecture",
            color=BODY_GRAY, size=12)
    p_role.add_run().add_break()
    add_run(p_role, "— AI & SaaS ISVs", color=BODY_GRAY, size=12)
    p_role.add_run().add_break()
    add_run(p_role, "Google Cloud", color=BODY_GRAY, size=12)

    p_spacer = content.add_paragraph()
    set_paragraph_spacing(p_spacer, before_pt=130, after_pt=0)

    p_date = content.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p_date, before_pt=0, after_pt=0)
    add_run(p_date, "May 2026", color=GRAY_MUTED, size=10)

    p_doctype = content.add_paragraph()
    p_doctype.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p_doctype, before_pt=2, after_pt=0)
    add_run(p_doctype, "ENGINEERING DESIGN DOCUMENT", color=COPPER,
            size=10, bold=True, char_spacing=2.0)


# ============================================================================
# Body section: headers, footers, page numbers
# ============================================================================

def configure_body_section(section):
    """Body section: 1in margins, running header + footer with page numbers."""
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    # Unlink header/footer from previous section so cover stays clean.
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    # Header: title on left, "Co-Innovation" on right (muted gray)
    header_p = section.header.paragraphs[0]
    set_paragraph_spacing(header_p, before_pt=0, after_pt=0)
    add_run(header_p, "Claude Powered Agents for Google Cloud  ",
            color=GRAY_MUTED, size=9)
    add_run(header_p, "•  Engineering Design Document",
            color=GRAY_MUTED, size=9)
    # Tab to right + co-innovation marker
    tab_run = header_p.add_run("\t")
    tab_run.font.name = FONT
    add_run(header_p, "CO-INNOVATION", color=GRAY_MUTED, size=9,
            char_spacing=1.5)
    # Bottom border on header for separation
    add_bottom_border(header_p, LIGHT_GRAY, size_eighths=4)
    # Configure tab stops: center-tab off, right-tab at 6.6"
    pPr = header_p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab1 = OxmlElement('w:tab')
    tab1.set(qn('w:val'), 'right')
    tab1.set(qn('w:pos'), str(int(6.6 * 1440)))
    tabs.append(tab1)
    pPr.append(tabs)

    # Footer: confidentiality left, "Page X" right
    footer_p = section.footer.paragraphs[0]
    set_paragraph_spacing(footer_p, before_pt=0, after_pt=0)
    add_run(footer_p, "Confidential — Engineering Reference",
            color=GRAY_MUTED, size=9)
    tab_run = footer_p.add_run("\t")
    tab_run.font.name = FONT
    add_run(footer_p, "Page  ", color=GRAY_MUTED, size=9)
    # PAGE field
    add_page_field(footer_p)
    pPr2 = footer_p._p.get_or_add_pPr()
    tabs2 = OxmlElement('w:tabs')
    tab2 = OxmlElement('w:tab')
    tab2.set(qn('w:val'), 'right')
    tab2.set(qn('w:pos'), str(int(6.6 * 1440)))
    tabs2.append(tab2)
    pPr2.append(tabs2)


def add_page_field(paragraph):
    """Insert a PAGE field that Word/Docs renders as the current page number."""
    run = paragraph.add_run()
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    for attr in ('ascii', 'hAnsi', 'cs', 'eastAsia'):
        rFonts.set(qn(f'w:{attr}'), FONT)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    rPr.append(sz)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), GRAY_MUTED)
    rPr.append(color)

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def restart_page_numbers(section):
    """Force the body section's page numbering to restart at 1."""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(existing)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)


# ============================================================================
# Table of contents (manual)
# ============================================================================

TOC_ENTRIES = [
    ("1.",  "Executive Summary",                                "5"),
    ("",    "1.1  What was built",                              "5"),
    ("",    "1.2  Why it matters — three audience lenses",      "5"),
    ("",    "1.3  The architectural thesis",                    "6"),
    ("",    "1.4  Why Option A wins over the alternatives",     "6"),
    ("",    "1.5  Status snapshot",                             "7"),
    ("2.",  "Use Case and Design Goals",                        "8"),
    ("",    "2.1  Problem statement",                           "8"),
    ("",    "2.2  The three Critical User Journeys",            "9"),
    ("",    "2.3  Design goals",                                "10"),
    ("",    "2.4  Non-goals",                                   "11"),
    ("3.",  "System Architecture",                              "12"),
    ("",    "3.1  Option A in prose and picture",               "12"),
    ("",    "3.2  Identity and trust flow",                     "13"),
    ("",    "3.3  Component map and responsibilities",          "14"),
    ("",    "3.4  Trust boundaries",                            "15"),
    ("4.",  "Component Deep-Dives",                             "16"),
    ("",    "4.1  The A2A Bridge (cc-a2a-bridge on Cloud Run)", "16"),
    ("",    "4.2  The ADK Orchestrator (cc-backend)",           "17"),
    ("",    "4.3  Claude Code as a Tool",                       "19"),
    ("",    "4.4  GKE Agent Sandbox (per-user pod)",            "20"),
    ("",    "4.5  Storage and persistence",                     "21"),
    ("",    "4.6  Artifact emission and MIME-aware routing",    "23"),
    ("",    "4.7  Workspace management tools",                  "25"),
    ("",    "4.8  Agent registration with Discovery Engine",    "26"),
    ("5.",  "A Request's Lifecycle",                            "27"),
    ("",    "5.1  Step-by-step walkthrough",                    "27"),
    ("",    "5.2  Parking and restoration",                     "29"),
    ("",    "5.3  Timing budget",                               "29"),
    ("6.",  "Deployment",                                       "30"),
    ("",    "6.1  GCP project layout",                          "30"),
    ("",    "6.2  Terraform structure",                         "31"),
    ("",    "6.3  Build pipeline",                              "32"),
    ("",    "6.4  Gated rollout pattern",                       "33"),
    ("",    "6.5  Deploying into a fresh GCP project",          "33"),
    ("7.",  "Best Practices and Deployment Learnings",          "34"),
    ("8.",  "Operational Considerations",                       "41"),
    ("9.",  "Extensibility and Roadmap",                        "44"),
    ("10.", "Appendix",                                         "47"),
]


def build_toc(doc):
    """Manual table of contents — two columns: title and page number."""
    add_h1(doc, "", "Table of Contents")
    # Override the h1: replace the leading ".  " with nothing
    # Simpler: write the heading ourselves
    last = doc.paragraphs[-1]
    last.clear()
    set_paragraph_spacing(last, before_pt=24, after_pt=18, line_pt=26)
    add_run(last, "Table of Contents", color=BLUE, size=20, bold=True)

    tbl = doc.add_table(rows=len(TOC_ENTRIES), cols=3)
    tbl.autofit = False
    tbl.allow_autofit = False
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9360')
    tblW.set(qn('w:type'), 'dxa')
    tbl._tbl.tblPr.append(tblW)
    remove_table_borders(tbl)

    for ri, (num, title, page) in enumerate(TOC_ENTRIES):
        row = tbl.rows[ri]
        c_num, c_title, c_page = row.cells
        set_cell_width(c_num, 0.35)
        set_cell_width(c_title, 5.4)
        set_cell_width(c_page, 0.7)
        for c in (c_num, c_title, c_page):
            set_cell_margins(c, top_in=0.03, bottom_in=0.03,
                             left_in=0.05, right_in=0.05)
            remove_cell_borders(c)

        is_section = num != ""
        size = 11 if is_section else 10
        bold = is_section
        color = NEAR_BLACK if is_section else BODY_GRAY

        p_num = c_num.paragraphs[0]
        set_paragraph_spacing(p_num, before_pt=0, after_pt=0, line_pt=15)
        if num:
            add_run(p_num, num, color=BLUE, size=size, bold=True)

        p_title = c_title.paragraphs[0]
        set_paragraph_spacing(p_title, before_pt=0, after_pt=0, line_pt=15)
        add_run(p_title, title, color=color, size=size, bold=bold)

        p_page = c_page.paragraphs[0]
        p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(p_page, before_pt=0, after_pt=0, line_pt=15)
        add_run(p_page, page, color=GRAY_MUTED, size=size)


# ============================================================================
# Body sections — content
# ============================================================================

def section_1_executive(doc):
    add_h1(doc, "1", "Executive Summary")

    add_h2(doc, "1.1", "What was built")

    add_para(doc,
        "This document describes Claude Powered Agents for Google Cloud, "
        "a first-party-quality entry in the Gemini Enterprise Agent "
        "Gallery that brings the Claude Code coding harness inside the "
        "customer's GCP project. A non-engineer — a product manager, "
        "growth marketer, designer, analyst, or legal counsel — selects "
        "Claude Code in any Gemini Enterprise thread, describes a goal "
        "in plain English, watches the agent read files, run commands, "
        "and write code in the working pane, and downloads the result "
        "as a file artifact in the same chat thread.")

    add_para(doc,
        "Inference runs on Vertex AI in the customer's tenant. Source "
        "data, scratch files, and generated artifacts live inside an "
        "ephemeral per-user workspace inside the customer's GKE "
        "cluster. Nothing leaves the project boundary. Identity flows "
        "from Gemini Enterprise through to per-user-prefix Cloud "
        "Storage access via downscoped STS tokens, enforced by IAM "
        "rather than by application code.")

    add_para(doc,
        "The agent has shipped in production, was registered with "
        "Gemini Enterprise Discovery Engine in Phase 7, and has been "
        "verified end-to-end against the three Critical User Journeys "
        "described in Section 2. As of May 2026 the agent is the live "
        "Claude entry in the Agent Gallery for the demo tenant, with "
        "downloads working, workspace state persistent across turns, "
        "and per-user storage isolation enforced by a passing negative "
        "test.")

    add_h2(doc, "1.2", "Why it matters — three audience lenses")

    add_para(doc,
        "For the customer, this is the path Claude Code into the "
        "enterprise without leaving the tenant. The non-engineer "
        "personas Anthropic's internal usage data identifies as the "
        "heaviest non-technical adopters — growth marketers, PMs, "
        "designers, analysts, finance, legal — are not going to install "
        "a CLI, configure Vertex credentials, or learn a development "
        "environment. They will, however, open Gemini Enterprise and "
        "describe what they want built. The agent meets them on that "
        "surface and produces a downloadable artifact. Data residency "
        "and IAM enforcement preserve the security posture the "
        "customer's CISO already approved.")

    add_para(doc,
        "For Anthropic, a Vertex-resident deployment of Claude Code "
        "demonstrates that Claude can be embedded as a peer of native "
        "Google Cloud services in regulated environments — not as a "
        "remote API customers must trust, but as a workload that runs "
        "inside the customer's project under the customer's IAM. "
        "Customers who would never approve outbound calls to "
        "api.anthropic.com routinely approve inbound calls from their "
        "Vertex AI projects to Anthropic-hosted endpoints, because "
        "Vertex is part of the existing privileged-API stack.")

    add_para(doc,
        "For Google Cloud, the Gemini Enterprise Agent Platform "
        "becomes a multi-model frontend rather than a single-model "
        "destination. The customer's users pick the best tool for the "
        "job from a gallery of specialist agents — Gemini for "
        "summarization, Claude for coding, third-party agents for "
        "domain-specific work — without giving up the Google identity, "
        "the audit log, the data-loss-prevention surface, or the "
        "single chat thread.")

    add_h2(doc, "1.3", "The architectural thesis")

    add_para(doc,
        "There were two plausible designs for this work. Option A is "
        "the one shipped: an Agent Development Kit (ADK) orchestrator "
        "agent runs inside the GKE pod and the Claude Code subprocess "
        "is one of its tools, alongside memory operations, artifact "
        "emission, and workspace management. The orchestrator is the "
        "brain; Claude Code is the executor for coding-shaped subtasks. "
        "Option B inverted this: Claude Code as the entry point, with "
        "ADK subagents underneath. The decision to ship Option A is "
        "the load-bearing architectural call in this build.")

    add_callout(doc, "WHY",
        "Option A keeps user identity, model choice, and tool surface "
        "orthogonal. Adding a second coding harness (Antigravity, "
        "OpenCode, Aider, Gemini's own code-execution tool) becomes a "
        "new tool registered on the same orchestrator, not a new "
        "architecture. Option B would have coupled all three together "
        "and forced a per-harness rewrite of identity propagation, "
        "memory, and artifact handling.")

    add_h2(doc, "1.4", "Why Option A wins over the alternatives")

    add_para(doc,
        "Option A separates concerns along the axis that matters for "
        "extension. Identity propagation, sandbox lifecycle, workspace "
        "persistence, signed-URL minting, memory, and artifact "
        "emission are all responsibilities of components that do not "
        "know which coding harness they're serving. The harness — "
        "Claude Code, in this build — is a subprocess invoked by a "
        "single tool with a stable contract: receive a prompt, "
        "stream events, read and write files under /workspace, "
        "return a final result. Replacing Claude Code with any other "
        "harness that honors that contract is a tool swap, not a "
        "platform change.")

    add_para(doc,
        "Option B, where Claude Code itself was the entry point and "
        "ADK sub-agents were invoked from inside it, would have "
        "required Claude Code to learn about Gemini Enterprise's "
        "A2A protocol, Cloud Run sandbox claims, downscoped STS "
        "tokens, and Discovery Engine registration semantics. Those "
        "concerns belong outside the harness. The Option-A separation "
        "also means the orchestrator can decide when the request "
        "doesn't need the coding harness at all — for greetings, "
        "follow-up questions, recall-only requests — and avoid "
        "spinning up a subprocess.")

    add_h2(doc, "1.5", "Status snapshot")

    add_table_styled(doc,
        headers=["Surface", "State", "Evidence"],
        rows=[
            ["Agent Gallery registration",
             "Live in the demo tenant",
             "Discovery Engine agent ID 5479509043993124503, "
             "state=ENABLED"],
            ["End-to-end CUJ coverage",
             "All three passing",
             "PM prototype, marketer CSV→ad-copy, analyst "
             "CSV→dashboard — verified in real GE threads"],
            ["Per-user isolation",
             "IAM-enforced, tested",
             "Negative test: user A's pod returns 403 when "
             "attempting to read user B's GCS prefix"],
            ["Workspace persistence",
             "Park/restore across pod cycles",
             "Manifest-driven backup to "
             "gs://<project>-cc-a2a-snapshots/users/<key>/, "
             "restored on next claim"],
            ["Multi-turn memory",
             "Cross-thread facts retained",
             "Firestore-backed `remember`/`recall` tools; "
             "facts survive pod restart and new GE thread"],
            ["Artifact downloads",
             "Both MIME paths verified",
             "Path A native chip for allowlisted MIMEs; Path B "
             "signed URL with Content-Disposition=attachment "
             "for the rest"],
        ],
        col_widths_in=[1.6, 1.7, 3.5],
    )


def section_2_use_case(doc):
    add_h1(doc, "2", "Use Case and Design Goals")

    add_h2(doc, "2.1", "Problem statement")

    add_para(doc,
        "The conventional Claude Code installation lives on an "
        "engineer's laptop or in a CI pipeline. It assumes a developer "
        "audience that knows what a CLI is, has cloud credentials, "
        "and edits files in an IDE. Anthropic's internal usage data "
        "tells a different story: by user count, the heaviest adopters "
        "of Claude Code are non-engineers in roles where the assistant "
        "writes single-shot artifacts — a marketer running A/B copy "
        "variants, a PM drafting a prototype to share with engineering, "
        "an analyst transforming a messy CSV into a dashboard, a legal "
        "counsel drafting a comparison table. They use Claude Code "
        "because the alternative tools require a coding intuition the "
        "task itself does not.")

    add_para(doc,
        "Those users are not going to install a developer environment, "
        "and the enterprise IT organization is rarely going to approve "
        "credentials on an unmanaged machine. What they will do — "
        "what they already do — is open Gemini Enterprise and chat "
        "with a gallery agent. The objective of this work is to put a "
        "Claude Code agent in that gallery so the chat thread becomes "
        "the surface on which the artifact is built.")

    add_para(doc,
        "The constraint that makes this hard is identity and "
        "isolation. Each Gemini Enterprise user must have a private "
        "workspace whose contents are scoped to them. Inference and "
        "stored data must remain inside the customer's GCP project. "
        "The CISO's existing audit trail, DLP scanning, and IAM "
        "configuration must remain authoritative — the agent cannot "
        "be a hole that bypasses them. And the experience inside the "
        "chat thread must be good enough that a non-engineer "
        "completes their task without leaving GE to find an "
        "alternative.")

    add_h2(doc, "2.2", "The three Critical User Journeys")

    add_para(doc,
        "Three Critical User Journeys (CUJs) were selected as the "
        "acceptance bar for v1. Each one exercises a distinct slice "
        "of the system, and together they cover the surface the agent "
        "must support for the documented user personas:")

    add_table_styled(doc,
        headers=["Persona", "Critical user journey",
                 "What it exercises"],
        rows=[
            ["Product manager",
             "Spec → prototype. Paste a PRD link; receive a "
             "working HTML prototype as a file artifact; iterate "
             "in the same thread (\"make the empty state friendlier\").",
             "Reading external content, generating multi-file HTML, "
             "Path A artifact emission, multi-turn refinement, "
             "workspace persistence across turns."],
            ["Growth marketer",
             "CSV → ad copy. Upload an ad-performance CSV; agent "
             "finds the bottom-quartile headlines and writes 50 "
             "variants; returns a spreadsheet artifact.",
             "Tabular data analysis, structured output to CSV, "
             "Path A allowlisted-MIME artifact emission, in-pod "
             "Python execution."],
            ["Analyst",
             "Messy CSV → interactive HTML dashboard with date "
             "slider. Drop a billing export; receive a self-contained "
             "HTML page with embedded JS and CSS.",
             "Data wrangling, complex HTML/JS generation, Path B "
             "signed-URL emission (HTML is not on the GE inline "
             "allowlist), Content-Disposition=attachment behaviour."],
        ],
        col_widths_in=[1.3, 2.9, 2.6],
    )

    add_para(doc,
        "These three were chosen because they cover the orthogonal "
        "dimensions of the agent's behavior: input modality (URL, "
        "uploaded CSV, dropped file), output modality (single-file "
        "artifact, multi-file artifact, large interactive artifact), "
        "and refinement pattern (one-shot, iterative). A system that "
        "ships all three is shaped right for the long tail of "
        "non-engineer tasks; a system that ships only one is "
        "load-bearing on a single happy path.")

    add_h2(doc, "2.3", "Design goals")

    add_h3(doc, "In-tenant data residency")

    add_para(doc,
        "All inference, all stored files, and all logs remain inside "
        "the customer's GCP project. Vertex AI is configured against "
        "the customer's project, not against a managed Anthropic "
        "endpoint. The workspace bucket lives in the customer's "
        "project. Firestore sessions and memory live in the customer's "
        "Firestore database. No outbound network calls to "
        "api.anthropic.com or to Anthropic's general-purpose hosted "
        "Claude. The Cloud Run bridge has a default-deny egress "
        "NetworkPolicy in front of the sandbox pods that prevents "
        "the harness from reaching the public internet at all.")

    add_h3(doc, "IAM-enforced per-user isolation")

    add_para(doc,
        "The agent must never let user A read user B's files. This "
        "is not enforced by application code that filters paths "
        "(which would fail under any vulnerability in the harness); "
        "it is enforced by IAM. Each user's pod receives a downscoped "
        "STS token with a Credential Access Boundary CEL condition "
        "that restricts the token's effective scope to "
        "gs://<bucket>/users/<user_key>/. The pod's own service "
        "account has zero IAM on the bucket. The negative test that "
        "demonstrates this — a pod with user_key A attempting to "
        "GET an object under users/B/ — receives a 403 from the "
        "Cloud Storage server, not a 403 from our code.")

    add_h3(doc, "Native Gemini Enterprise chat surface")

    add_para(doc,
        "The agent's UI is the Gemini Enterprise thread. There is no "
        "separate web app, no installed client, no embedded iframe. "
        "Tool calls render in the thinking pane as the agent runs. "
        "File artifacts appear as native GE chips when MIME is "
        "supported. The pattern matches what a Gemini Enterprise "
        "user already knows; the Claude attribution is in the agent "
        "card and the working pane, not in a divergent UX.")

    add_h3(doc, "Downloadable artifacts in the chat thread")

    add_para(doc,
        "Every file the agent produces must be downloadable directly "
        "from the chat thread without context-switching. The agent "
        "emits artifacts via the emit_artifact tool, which routes "
        "automatically based on MIME: allowlisted types render as "
        "native chips; everything else returns a signed Cloud "
        "Storage URL with Content-Disposition=attachment that the "
        "agent embeds in its reply as a clickable link.")

    add_h3(doc, "Conversational workspace management")

    add_para(doc,
        "The workspace is not a hidden implementation detail. The "
        "user can ask \"what files do I have?\", \"open the dashboard "
        "I made yesterday\", \"delete the old draft\", and the agent "
        "responds without invoking the Claude Code subprocess. The "
        "list_workspace, read_workspace_file, delete_workspace_file, "
        "and move_workspace_file tools live on the orchestrator and "
        "respond directly. The workspace is treated as the user's "
        "filesystem; the agent is its concierge.")

    add_h2(doc, "2.4", "Non-goals")

    add_para(doc,
        "Several features that would be reasonable in a longer-term "
        "version are explicitly out of scope for v1, so the surface "
        "stays small enough to verify end-to-end. The agent is not a "
        "general-purpose IDE replacement; it does not provide "
        "syntax highlighting, debugger surfaces, or git-history "
        "browsing inside the chat thread. It is not a "
        "multi-user collaborative workspace; each user's pod is "
        "private, and shared state across users is deferred to the "
        "memory layer and explicit handoff patterns. It does not "
        "embed model-routing logic that picks between Claude variants "
        "for different subtasks; v1 uses Claude Opus 4.7 for both "
        "the orchestrator (via LiteLLM/ADK) and the harness (via "
        "Vertex). It does not aim to support coding harnesses other "
        "than Claude Code in v1, although the architecture admits "
        "additional harnesses as new tools without structural change.")


def section_3_architecture(doc):
    add_h1(doc, "3", "System Architecture")

    add_h2(doc, "3.1", "Option A in prose and picture")

    add_para(doc,
        "The system has three layers, each of which is the boundary "
        "of a different concern. The Gemini Enterprise frontend is "
        "the user's surface and the source of identity. The A2A "
        "bridge on Cloud Run is the stateless protocol adapter that "
        "translates between Gemini Enterprise's Agent-to-Agent "
        "protocol and the in-cluster sandbox runtime. The cc-backend "
        "pod, running in a GKE Agent Sandbox, hosts the ADK "
        "orchestrator and the Claude Code subprocess for a single "
        "user. The bridge is stateless and shared; the backend pod "
        "is stateful and per-user.")

    add_code_block(doc,
        "Gemini Enterprise (A2A turn + user OAuth token)\n"
        "        │\n"
        "        ▼\n"
        "┌────────────────────────────────────────────────┐\n"
        "│  cc-a2a-bridge  (Cloud Run, stateless)         │\n"
        "│  - A2A JSON-RPC adapter only                   │\n"
        "│  - Resolves token → user_key                   │\n"
        "│  - Gets-or-creates SandboxClaim/cc-u-<key>     │\n"
        "│  - Mints downscoped STS token for user prefix  │\n"
        "│  - Streams events ⇄ A2A working/thought/result │\n"
        "└────────────────┬───────────────────────────────┘\n"
        "                 │ (via internal LB, X-Sandbox-* headers)\n"
        "                 ▼\n"
        "┌────────────────────────────────────────────────┐\n"
        "│  cc-backend  (GKE Agent Sandbox, per-user pod) │\n"
        "│  - gVisor-isolated, default-deny networking    │\n"
        "│  - /workspace = 20Gi ephemeral, noexec         │\n"
        "│  - ADK agent (Python) — the orchestrator       │\n"
        "│      ├── claude_code tool (drives CLI subproc) │\n"
        "│      ├── workspace tools (list/read/move/del)  │\n"
        "│      ├── memory tools (remember/recall)        │\n"
        "│      └── emit_artifact tool                    │\n"
        "│  - One ADK Session per A2A context_id          │\n"
        "│  - Park/restore /workspace to GCS              │\n"
        "└────────────────────────────────────────────────┘"
    )

    add_para(doc,
        "Each A2A request from Gemini Enterprise carries an "
        "Authorization header containing a Google access token "
        "minted on behalf of the end user. The bridge resolves the "
        "token through oauth2.googleapis.com/tokeninfo, extracts the "
        "user's email or subject identifier, hashes it to produce a "
        "short user_key, and uses that key as the suffix for the "
        "SandboxClaim resource it gets-or-creates against the GKE "
        "Agent Sandbox controller. The hashing ensures user PII does "
        "not appear in K8s resource names or logs.")

    add_h2(doc, "3.2", "Identity and trust flow")

    add_para(doc,
        "Identity flows top-down through the system and is checked "
        "or downgraded at each boundary. Gemini Enterprise mints the "
        "user OAuth token and routes the A2A request to the bridge. "
        "The bridge resolves the token (one round-trip to "
        "tokeninfo, cached per user for the duration of the bridge "
        "instance) and computes user_key. The bridge does not pass "
        "the user's OAuth token to the backend; instead, it mints "
        "two derived credentials. First, a service-to-service "
        "identity token for the cc-backend pod (so the pod can "
        "verify the request came from the trusted bridge). Second, "
        "a downscoped STS token, scoped via Credential Access "
        "Boundary CEL to users/<user_key>/ inside the workspace "
        "bucket, that the backend uses for any GCS access.")

    add_callout(doc, "INVARIANT",
        "The backend pod's service account has zero IAM on the "
        "workspace bucket. Every GCS operation the backend performs "
        "uses the downscoped STS token minted by the bridge. If the "
        "bridge stops minting that token, the backend cannot reach "
        "GCS at all. This is the load-bearing isolation guarantee "
        "and the property tested by the negative test.")

    add_para(doc,
        "Inside the pod, the ADK orchestrator decides which tool to "
        "run for a given user turn. When it invokes the Claude Code "
        "subprocess, it passes /workspace as the working directory "
        "and the downscoped credentials as environment. The Claude "
        "Code CLI itself does not authenticate to GCS directly; "
        "anything it writes goes to /workspace, and the orchestrator's "
        "emit_artifact tool is responsible for moving files from "
        "/workspace to GCS or for synthesizing a signed URL.")

    add_h2(doc, "3.3", "Component map and responsibilities")

    add_table_styled(doc,
        headers=["Component", "Where it runs", "Owns"],
        rows=[
            ["Gemini Enterprise UI",
             "Customer's GE tenant",
             "User chat surface, identity, agent gallery"],
            ["Discovery Engine",
             "Customer's GE tenant",
             "Agent card lookup, A2A request routing"],
            ["cc-a2a-bridge",
             "Cloud Run, customer project, us-central1",
             "A2A protocol adapter, token resolution, "
             "SandboxClaim lifecycle, signed-URL minting"],
            ["sandbox-router",
             "GKE, in-cluster, cc-sandbox namespace",
             "X-Sandbox-* header routing to per-user pod IPs"],
            ["cc-backend (per-user pod)",
             "GKE Agent Sandbox, gVisor runtime",
             "ADK orchestrator, Claude Code subprocess, "
             "/workspace, park/restore"],
            ["Workspace bucket",
             "GCS, customer project",
             "User workspace persistence, signed URL targets"],
            ["Firestore (cc-on-ge database)",
             "Customer project",
             "ADK Sessions per context_id, long-term memory per "
             "user_key"],
            ["Vertex AI (Claude)",
             "Customer project, global endpoint",
             "Inference for both orchestrator and Claude Code "
             "subprocess"],
        ],
        col_widths_in=[1.7, 2.0, 3.1],
    )

    add_h2(doc, "3.4", "Trust boundaries")

    add_para(doc,
        "Four trust boundaries matter. The first is Gemini Enterprise "
        "to the bridge: GE authenticates as the user, the bridge "
        "trusts GE's OAuth issuer chain. The second is the bridge "
        "to the per-user pod: the bridge is the only caller the pod "
        "accepts (via Cloud Run identity tokens), and the bridge is "
        "the only entity that can claim or release a pod. The third "
        "is the per-user pod to Cloud Storage: enforced by the "
        "downscoped STS token. The fourth is the per-user pod to "
        "the rest of the internet, which is closed by a default-deny "
        "NetworkPolicy on the cc-sandbox namespace. Egress is "
        "permitted only to the Vertex AI endpoint, the internal "
        "router service, and the bridge's signing endpoint.")


def section_4_components(doc):
    add_h1(doc, "4", "Component Deep-Dives")

    # ---- 4.1 Bridge ----
    add_h2(doc, "4.1", "The A2A Bridge (cc-a2a-bridge on Cloud Run)")

    add_para(doc,
        "The bridge is a small Python FastAPI service that "
        "implements the Agent-to-Agent (A2A) JSON-RPC protocol "
        "consumed by Gemini Enterprise Discovery Engine. It has no "
        "LLM logic and no per-user state; everything stateful lives "
        "in the backend pod. The bridge is deployed on Cloud Run "
        "for three reasons: it scales independently of the GKE "
        "cluster (a hot bridge does not need warm pods), it benefits "
        "from Cloud Run's built-in HTTPS termination and identity "
        "verification, and keeping it out of the cluster avoids a "
        "circular-dependency between the bridge image and the "
        "Sandbox-Template image during deploys.")

    add_h3(doc, "A2A protocol adapter responsibilities")

    add_para(doc,
        "On every incoming request, the bridge implements the A2A "
        "verbs message/send and message/stream. For message/stream "
        "it opens an SSE upstream to the backend's /execute endpoint "
        "and translates the backend's event stream into A2A "
        "working/thought/response events. Tool calls become "
        "structured A2A artifact events; thinking blocks become "
        "collapsed thought blocks; files become A2A FileWithBytes "
        "(Path A) or text with embedded signed URLs (Path B).")

    add_h3(doc, "The /workspace/sign endpoint")

    add_para(doc,
        "The bridge exposes an internal /workspace/sign endpoint "
        "that the backend calls when it needs to produce a "
        "downloadable URL for a file that cannot be inlined in an "
        "A2A FileWithBytes (because its MIME is not on the verified "
        "GE allowlist, or because the file exceeds 5 MB). The "
        "endpoint requires the X-Sandbox-User-Key header that the "
        "router enforces is bound to the calling pod's identity. "
        "It mints a v4 GCS signed URL with response_disposition="
        "attachment to force download behavior in the browser.")

    add_callout(doc, "PHASE 1 INVARIANT",
        "The bridge is the only component in the system that holds "
        "GCS signing credentials. The backend pod's service account "
        "has zero IAM on the workspace bucket. This invariant is "
        "what allows the entire isolation guarantee to be checked "
        "by inspecting one component's IAM bindings — there is "
        "exactly one signing surface to audit.")

    add_para(doc,
        "The signing path was hardened in May 2026 by adding "
        "Content-Disposition=attachment with a filename derived "
        "from the workspace-relative path, sanitized to "
        "[A-Za-z0-9._-] to close the header-injection surface. "
        "Before that fix, signed URLs to HTML files rendered "
        "inline in the browser, and the agent had emitted "
        "right-click apology language as a workaround. The fix is "
        "documented in bridge/sign_helpers.py with a 12-case "
        "regression test suite.")

    add_h3(doc, "IAM model")

    add_para(doc,
        "The bridge's Cloud Run service has no public invoker "
        "binding. roles/run.invoker on the bridge service is "
        "granted only to two principals: the cc-backend service "
        "account (so the pod can call /workspace/sign), and the "
        "Discovery Engine service agent for the customer project "
        "(so GE can route A2A requests to the bridge). "
        "Unauthenticated access from any other source is rejected "
        "at the Cloud Run frontend before reaching the application.")

    # ---- 4.2 ADK Orchestrator ----
    add_h2(doc, "4.2", "The ADK Orchestrator (cc-backend)")

    add_para(doc,
        "The orchestrator is a Google Agent Development Kit (ADK) "
        "agent — a Python object whose definition is a name, a "
        "model identifier, a tool registry, and a system prompt. "
        "It runs once per A2A context_id, with state persisted to "
        "a custom FirestoreSessionService implementation that "
        "stores the ADK event log in the cc-on-ge Firestore "
        "database under sessions/<context_id>. On each user turn, "
        "ADK replays the event log, asks the orchestrator to "
        "produce a response, and the orchestrator decides — via "
        "the model — which tools to call.")

    add_h3(doc, "Why ADK rather than raw Claude Code at the entry point")

    add_para(doc,
        "ADK gives us multi-turn state, a typed tool registry with "
        "automatic JSON schema generation for function calling, "
        "and a session model that aligns naturally with A2A "
        "context_ids. Reimplementing those concerns inside Claude "
        "Code would have required intrusive changes to the harness "
        "and would have coupled the platform to Claude. The "
        "orchestrator-as-ADK choice keeps Claude Code as a clean "
        "subprocess invoked by one tool.")

    add_h3(doc, "The orchestrator's tool registry")

    add_table_styled(doc,
        headers=["Tool", "Responsibility"],
        rows=[
            ["claude_code",
             "Invoke the Claude Code CLI as a subprocess against "
             "/workspace; stream tool-call events back to the "
             "orchestrator; return the final result text."],
            ["emit_artifact",
             "Move a file from /workspace into the A2A response. "
             "Routes by MIME: allowlisted → Path A (FileWithBytes); "
             "non-allowlisted or large → Path B (signed URL)."],
            ["list_workspace",
             "Return a one-shot listing of /workspace contents — "
             "filenames, sizes, modification times — without "
             "invoking the Claude Code harness."],
            ["read_workspace_file",
             "Read a single file from /workspace and return its "
             "content inline. Used for follow-up questions like "
             "\"what was in that script you wrote?\""],
            ["delete_workspace_file",
             "Two-step soft delete: first call moves the file to "
             "/workspace/.trash/, second call (with confirm=True) "
             "hard-deletes. Prevents single-fat-finger data loss."],
            ["move_workspace_file",
             "Rename or relocate a file within /workspace, "
             "including into and out of .trash for undo."],
            ["get_download_url",
             "Produce a Path B signed URL for an existing file in "
             "/workspace, without re-emitting it as an artifact."],
            ["remember",
             "Write a fact to long-term memory under "
             "memory/<user_key>/facts/. Used for cross-thread "
             "preferences and project context."],
            ["recall",
             "Read facts from long-term memory. Invoked at the "
             "start of each turn so the orchestrator can ground "
             "the response in what the user has previously told "
             "it."],
        ],
        col_widths_in=[1.7, 5.1],
    )

    add_h3(doc, "System prompt structure")

    add_para(doc,
        "The orchestrator's system prompt teaches it three things "
        "in order: that recall is the first action of every turn so "
        "long-term context lands in the conversation before the "
        "model decides what to do; that the workspace tools should "
        "be preferred over claude_code for any question about "
        "existing files; and that emit_artifact must be called in "
        "the same turn as the file write, with Path A and Path B "
        "templates supplied for how to phrase the response. The "
        "prompt explicitly suppresses workaround language "
        "(\"right-click and save link as\", \"your browser may "
        "preview it as text\") that emerged as a model artifact "
        "when the Content-Disposition fix was outstanding.")

    add_h3(doc, "Multi-turn behavior within a pod")

    add_para(doc,
        "An ADK Session corresponds to one Gemini Enterprise "
        "context_id, which corresponds to one chat thread. The "
        "Session's event log is persisted to Firestore on every "
        "tool call so that a pod restart mid-session resumes "
        "correctly. A second user turn on the same context_id "
        "lands on the same pod (the SandboxClaim is sticky) and "
        "replays the prior session state. A user with multiple "
        "open GE threads has multiple Sessions in their pod, all "
        "sharing the same workspace.")

    # ---- 4.3 Claude Code as Tool ----
    add_h2(doc, "4.3", "Claude Code as a Tool")

    add_para(doc,
        "The claude_code tool wraps the claude-agent-sdk Python "
        "client, which is the supported way to drive the Claude "
        "Code CLI from Python. The tool accepts a prompt argument "
        "and an optional working_directory (defaulted to "
        "/workspace), spawns the CLI as a subprocess with the "
        "appropriate environment, and streams tool-call events back "
        "to the orchestrator as they arrive. The orchestrator "
        "surfaces those events to the A2A bridge for rendering in "
        "the Gemini Enterprise working pane.")

    add_h3(doc, "bypassPermissions mode and IS_SANDBOX=1")

    add_para(doc,
        "The Claude Code CLI normally prompts for permission on "
        "destructive operations. Inside the per-user sandbox, "
        "permission prompts have no useful UI surface — they would "
        "block the agent on a question the user cannot answer from "
        "the chat thread. The pod sets IS_SANDBOX=1 (signaling to "
        "the harness that it is running in an isolated environment "
        "with no shared state) and invokes the CLI with "
        "--dangerously-skip-permissions. The blast radius is "
        "bounded by the sandbox itself: gVisor, default-deny "
        "network, /workspace mounted noexec, no host access. "
        "Anything destructive Claude Code does is destructive only "
        "to that single user's ephemeral workspace.")

    add_h3(doc, "Allowed tools")

    add_para(doc,
        "The tool exposes the Claude Code CLI with Read, Write, "
        "Edit, Bash, Grep, and Glob enabled. The MCP and Agent "
        "tools are disabled — the orchestrator is responsible for "
        "agent-shaped behavior, and MCP server registration would "
        "create a side-channel through which a harness invocation "
        "could affect another user. WebFetch and WebSearch are "
        "disabled because the egress NetworkPolicy already closes "
        "the public-internet surface and enabling them would just "
        "produce failing tool calls.")

    add_h3(doc, "Per-turn subprocess isolation")

    add_para(doc,
        "Each invocation of the claude_code tool is a fresh "
        "subprocess. The subprocess inherits /workspace and the "
        "environment, but no in-memory state from prior runs. This "
        "is intentional: it means the orchestrator's view of the "
        "world (the ADK Session, the memory facts, the workspace "
        "listing) is the single source of truth for what the agent "
        "knows. The harness is stateless; the orchestrator is "
        "stateful.")

    # ---- 4.4 GKE Sandbox ----
    add_h2(doc, "4.4", "GKE Agent Sandbox (per-user pod)")

    add_para(doc,
        "The pod runs under the GKE Agent Sandbox addon "
        "(extensions.agents.x-k8s.io API group), which gives us a "
        "first-class CRD model for per-user-pod lifecycle. A "
        "SandboxTemplate names the pod spec; a SandboxClaim, "
        "created by the bridge with name cc-u-<user_key>, "
        "instantiates a pod from that template if one does not "
        "already exist. A WarmPool resource keeps a small set of "
        "templated pods pre-spawned and ready to be bound to "
        "incoming claims, so the first turn for a returning user "
        "feels instant rather than spending 8-15 seconds on cold "
        "container start.")

    add_h3(doc, "Sandbox isolation primitives")

    add_para(doc,
        "Three primitives stack to produce the isolation guarantee. "
        "First, runtimeClassName: gvisor — the pod's containers "
        "execute under gVisor's user-space kernel, which interposes "
        "on syscalls and blocks the classes of escape historically "
        "used in container escapes. Second, /workspace is an "
        "ephemeral PVC mounted with noexec — files written to "
        "/workspace cannot be executed directly, which prevents a "
        "compromised harness from staging a binary in the workspace "
        "and running it. Third, the cc-sandbox namespace has a "
        "default-deny NetworkPolicy plus explicit allow rules for "
        "the router service and the Vertex endpoint; sandbox-to-"
        "sandbox network traffic is closed.")

    add_h3(doc, "Resource limits and lifecycle")

    add_para(doc,
        "Each pod requests 2 vCPU and 4Gi memory, with the "
        "20Gi /workspace PVC sized to accommodate the largest "
        "expected artifact bundle. Pods are released by the bridge "
        "30 minutes after the most recent A2A request, on the "
        "assumption that idle workspaces should not keep node "
        "capacity reserved. The sweeper job in the bridge "
        "iterates active SandboxClaims every 5 minutes and triggers "
        "park-and-release for any whose last-activity timestamp "
        "exceeds the idle threshold.")

    add_h3(doc, "The warm pool")

    add_para(doc,
        "The WarmPool resource configures a pool of two pre-spawned "
        "pods that are bound to claims on demand. Two is enough for "
        "the demo tenant; the pool size is a tuning knob that "
        "trades node cost for tail latency on first-turn requests. "
        "A pod that has been claimed and released does not return "
        "to the warm pool — it is destroyed, and the WarmPool "
        "controller spawns a replacement against the current "
        "SandboxTemplate. This means template updates roll forward "
        "naturally as claims cycle, without an explicit pod recycle.")

    # ---- 4.5 Storage and persistence ----
    add_h2(doc, "4.5", "Storage and persistence")

    add_h3(doc, "GCS layout")

    add_para(doc,
        "The workspace bucket is gs://<project>-cc-a2a-snapshots, "
        "and every user's data lives under "
        "gs://<project>-cc-a2a-snapshots/users/<user_key>/. Inside "
        "a user's prefix, the layout is a tarball-style snapshot "
        "named with the SandboxClaim and the timestamp, plus a "
        "current/manifest.json that lists files and SHAs so a "
        "subsequent claim can quickly determine what to restore.")

    add_h3(doc, "Park and restore")

    add_para(doc,
        "When a SandboxClaim is released, the bridge issues a POST "
        "to the pod's /park endpoint, which serializes /workspace "
        "into a tarball, uploads it to the user's prefix, and "
        "writes the manifest. When a fresh claim binds (either "
        "first time or after release), the pod's startup probe "
        "fetches the most recent manifest and restores files into "
        "/workspace before accepting requests. The manifest format "
        "is chosen deliberately to make incremental sync possible "
        "in a future iteration without changing the on-disk format.")

    add_h3(doc, "Firestore: sessions and memory")

    add_para(doc,
        "Two Firestore collections back the agent's state. "
        "sessions/{context_id} stores the ADK Session for a given "
        "Gemini Enterprise thread: the event log, including every "
        "tool call and every model response, that ADK replays on "
        "the next turn. memory/{user_key}/facts/{fact_id} stores "
        "long-term, cross-thread facts: things the user has asked "
        "the agent to remember, such as preferred output formats, "
        "names of recurring projects, or stylistic preferences. "
        "The orchestrator's remember and recall tools are the "
        "interface to this collection.")

    add_h3(doc, "Downscoped STS tokens")

    add_para(doc,
        "The bridge holds a service account (cc-a2a-bridge) with "
        "roles/storage.objectAdmin on the workspace bucket. For "
        "every backend request, it mints a downscoped STS token "
        "via google.auth.downscoped.Credentials, with a "
        "CredentialAccessBoundary specifying a single "
        "AccessBoundaryRule whose AvailabilityCondition is the CEL "
        "expression "
        "resource.name.startsWith('projects/_/buckets/<bucket>/"
        "objects/users/<user_key>/'). The resulting token has at "
        "most read/write access to that prefix and is forwarded to "
        "the pod via an X-Sandbox-Token header that the router "
        "preserves. Tokens are short-lived (30 minutes) and minted "
        "fresh for every claim binding.")

    add_callout(doc, "EVIDENCE",
        "The negative test in make iso-test creates two pods with "
        "distinct user_keys and verifies that pod A, when given "
        "its downscoped token and asked to list users/B/, receives "
        "a 403 from storage.googleapis.com — not a 403 from our "
        "code. This is the externally-observable signal that the "
        "STS Credential Access Boundary is doing the work, not an "
        "application-layer filter.")

    # ---- 4.6 Artifacts ----
    add_h2(doc, "4.6", "Artifact emission and MIME-aware routing")

    add_para(doc,
        "Files the agent produces under /workspace become "
        "downloadable artifacts via the emit_artifact tool. The "
        "tool's responsibility is to make the file reachable from "
        "the chat thread; the question is whether to inline it "
        "(Path A) or to mint a signed URL the agent embeds in its "
        "reply (Path B). The decision is made by MIME type and "
        "size, and is the single point at which the agent's "
        "presentation behavior is configured.")

    add_h3(doc, "Path A: native FileWithBytes chip")

    add_para(doc,
        "If the MIME type is on the verified Gemini Enterprise "
        "allowlist and the file is under 5 MB, emit_artifact "
        "reads the file from /workspace, base64-encodes it into an "
        "A2A FileWithBytes structure, and returns it as part of "
        "the response. The bridge translates the FileWithBytes "
        "into an A2A artifact event; GE renders it as a native "
        "download chip in the thread, with the agent's text reply "
        "above. The user clicks the chip; GE serves the file "
        "directly to the browser. No external storage involvement.")

    add_h3(doc, "Path B: signed Cloud Storage URL")

    add_para(doc,
        "If the MIME is not allowlisted, or the file exceeds the "
        "5 MB inline ceiling, the tool first uploads the file to "
        "gs://<bucket>/users/<user_key>/out/<filename> using the "
        "downscoped STS token, then asks the bridge to mint a v4 "
        "signed GET URL with response_disposition=attachment and "
        "filename=<sanitized basename>. The URL is returned to "
        "the orchestrator, which embeds it in the agent's reply as "
        "a clickable link. The URL is valid for 15 minutes; the "
        "user clicks; the browser triggers a download.")

    add_h3(doc, "Why \"everything else → Path B\" is the right default")

    add_para(doc,
        "Gemini Enterprise's UI renders native chips only for a "
        "small set of MIME types — primarily image/png, image/jpeg, "
        "application/pdf, text/csv, text/plain, application/json, "
        "and a handful of Office document MIMEs. For everything "
        "else, including text/html, the chip is rejected and the "
        "user sees nothing. Path B exists because Path A's set is "
        "narrower than the agent's output surface. The decision to "
        "make Path B the universal fallback — rather than tailoring "
        "per-type fallback strategies — keeps the agent's "
        "behavioral surface small and testable.")

    add_h3(doc, "Content-Disposition=attachment (May 2026 hotfix)")

    add_para(doc,
        "An earlier version of the signing path produced URLs "
        "without response_disposition. For text/html files, the "
        "browser interpreted the Content-Type header as authoritative "
        "and rendered the page inline rather than downloading it. "
        "The agent had been emitting workaround language "
        "(\"right-click → save link as\") to compensate for the "
        "broken UX. The May 2026 fix adds attachment behavior to "
        "every signed URL and sanitizes the filename to "
        "[A-Za-z0-9._-] to close the header-injection surface in "
        "the Content-Disposition value. The 12-case regression "
        "suite in bridge/test_sign.py covers happy path, header "
        "injection attempts, degenerate names ('', '.', '..'), and "
        "the 200-character truncation that preserves extensions.")

    # ---- 4.7 Workspace tools ----
    add_h2(doc, "4.7", "Workspace management tools")

    add_para(doc,
        "The workspace is not opaque. The user can ask about it "
        "directly, and the orchestrator answers without invoking "
        "the Claude Code subprocess. Four tools — list_workspace, "
        "read_workspace_file, delete_workspace_file, and "
        "move_workspace_file — implement the conversational "
        "workspace surface, with a fifth tool, get_download_url, "
        "for the case where the user wants to re-download a file "
        "they had emitted earlier in the thread.")

    add_h3(doc, "Two-step soft-delete")

    add_para(doc,
        "delete_workspace_file is intentionally a two-step "
        "operation. The first call (confirm=False, the default) "
        "moves the file from its current path to "
        "/workspace/.trash/<basename>-<timestamp> and reports back "
        "to the user. The second call (with confirm=True) hard-"
        "deletes from /workspace/.trash. The pattern exists because "
        "the harness occasionally generates a delete request "
        "ambiguously — a user saying \"clean up that draft\" can "
        "be reasonably interpreted as either soft or hard delete — "
        "and the soft-delete is a recoverable single-fat-finger "
        "boundary. move_workspace_file admits moves into and out "
        "of .trash, which gives the user an undo path.")

    add_h3(doc, "Function-calling argument coercion")

    add_para(doc,
        "Tool arguments arrive at the workspace tools as strings, "
        "even when the schema declares them as integers, booleans, "
        "or lists. This is a property of how function calling "
        "lowers structured tool calls onto the wire — typed schemas "
        "describe the contract, but the values are stringified for "
        "transport. An earlier version of read_workspace_file took "
        "an offset parameter typed as int and crashed when the "
        "model passed \"0\". A May 2026 hotfix added explicit "
        "coercion at the tool boundary (int(arg), bool(arg) with "
        "string truthiness), backed by a 15-case regression suite. "
        "The lesson generalizes: typed arguments coming from "
        "function calling should be coerced at the boundary, not "
        "trusted to arrive in their declared type.")

    # ---- 4.8 Registration ----
    add_h2(doc, "4.8", "Agent registration with Discovery Engine")

    add_para(doc,
        "The agent is registered with Gemini Enterprise via "
        "Discovery Engine's Agents API. Registration is a one-time "
        "operation per environment, performed by the scripts/register-"
        "agent.sh helper, that creates an Agent resource referencing "
        "the bridge's A2A endpoint URL and the agent card published "
        "at /.well-known/agent-card.json on the bridge.")

    add_h3(doc, "The agent card")

    add_para(doc,
        "The agent card is a JSON document that describes the "
        "agent to GE: name, description, supported A2A skills, "
        "input/output modalities, and any optional authorization "
        "configuration. The card declares the agent supports "
        "message/send and message/stream, that responses may "
        "include FileWithBytes artifacts, and that the input is "
        "free-form text plus optional file uploads. The card is "
        "served by the bridge at a public well-known path, but "
        "the underlying service requires Cloud Run invoker IAM, "
        "so only Discovery Engine's service agent can actually "
        "load it.")

    add_h3(doc, "Authorization")

    add_para(doc,
        "Version 1 ships in single-user demo mode: the bridge "
        "trusts the OAuth token GE forwards and uses tokeninfo "
        "lookup to derive identity. The per-user authorizationConfig "
        "surface that Discovery Engine provides — which would "
        "route a customer's GE users through per-user OAuth "
        "consent screens to mint application-specific tokens — is "
        "deferred to v2. The deferral is documented in PROJECT_PLAN.md "
        "as a Phase 7 prep limitation and is bounded by the fact "
        "that the bridge already supports user_key derivation and "
        "the rest of the system is keyed off user_key.")


def section_5_lifecycle(doc):
    add_h1(doc, "5", "A Request's Lifecycle")

    add_h2(doc, "5.1", "Step-by-step walkthrough")

    add_para(doc,
        "A single user turn proceeds through eleven distinct hops, "
        "each with its own identity and trust posture. Tracing one "
        "request end-to-end makes the system's behavior concrete.")

    add_para(doc,
        "(1) A user types a request in a Gemini Enterprise thread. "
        "GE generates an A2A request — a JSON-RPC message/stream "
        "call — and includes an Authorization: Bearer <token> "
        "header where the token is a Google OAuth access token "
        "minted on behalf of the end user. The request's "
        "context_id is the GE thread identifier; the task_id is "
        "unique per turn.")

    add_para(doc,
        "(2) Discovery Engine resolves the agent card, which lists "
        "the bridge URL, and forwards the A2A request to "
        "cc-a2a-bridge on Cloud Run. The forwarding identity is "
        "the Discovery Engine service agent for the customer's "
        "project — that identity has roles/run.invoker on the "
        "bridge service, so the request is admitted past Cloud "
        "Run's IAM check.")

    add_para(doc,
        "(3) The bridge resolves the user's OAuth token via "
        "oauth2.googleapis.com/tokeninfo, extracts the email or "
        "subject, and derives user_key by SHA256-truncate. It "
        "logs the resolution at INFO level (the user_key, not the "
        "email — PII never appears in bridge logs).")

    add_para(doc,
        "(4) The bridge calls the K8s API to GET the SandboxClaim "
        "cc-u-<user_key> in the cc-sandbox namespace. If the "
        "claim exists and has a bound pod, the bridge reuses it. "
        "If not, the bridge creates the claim, which causes the "
        "Agent Sandbox controller to bind a pod from the warm "
        "pool. The bound pod's IP becomes the routing target.")

    add_para(doc,
        "(5) The bridge mints a downscoped STS token scoped to "
        "users/<user_key>/, formats a backend request including "
        "the prompt, the context_id, and the X-Sandbox-Token "
        "header, and POSTs to the in-cluster router service.")

    add_para(doc,
        "(6) The router examines the X-Sandbox-* headers and "
        "forwards the request to the bound pod's port 9000. The "
        "router is a simple TCP/HTTP forwarder; its only routing "
        "policy is the header-to-pod mapping derived from "
        "SandboxClaim state.")

    add_para(doc,
        "(7) The pod's /execute handler receives the request, "
        "looks up or instantiates the ADK Session for the "
        "context_id (Firestore read), and invokes the ADK "
        "orchestrator with the user's prompt as the new turn.")

    add_para(doc,
        "(8) The orchestrator runs through its tool-call loop. "
        "The first action is typically recall against the user's "
        "long-term memory. Depending on the request, the next "
        "actions are some combination of list_workspace, "
        "read_workspace_file, or claude_code. Each tool call is "
        "streamed back to the bridge over SSE.")

    add_para(doc,
        "(9) When the orchestrator emits artifacts, emit_artifact "
        "routes by MIME. Path A inlines bytes; Path B uploads to "
        "GCS via the downscoped STS token and asks the bridge to "
        "mint a signed URL via /workspace/sign.")

    add_para(doc,
        "(10) The bridge translates the backend's SSE event stream "
        "into A2A working/thought/response events, "
        "re-serializes file artifacts as FileWithBytes (Path A) "
        "or text-with-embedded-URL (Path B), and forwards them to "
        "Discovery Engine.")

    add_para(doc,
        "(11) Discovery Engine forwards the events to the user's "
        "GE thread, where they render in the working pane (tool "
        "calls), the thinking pane (collapsed thought blocks), "
        "and the response pane (text plus chips or links). The "
        "user sees the agent's response and downloads artifacts "
        "directly from the chat.")

    add_h2(doc, "5.2", "Parking and restoration")

    add_para(doc,
        "When the bridge's sweeper detects that a SandboxClaim "
        "has been idle for 30 minutes, it issues POST /park to "
        "the pod. The pod responds by tarring /workspace into a "
        "snapshot, uploading it under the user's GCS prefix, "
        "writing a manifest.json with file sizes and SHAs, and "
        "deleting itself. The next claim for the same user_key "
        "binds a fresh pod, which on startup reads "
        "current/manifest.json from GCS and restores /workspace "
        "before accepting any /execute calls. The restoration is "
        "transparent to the user — their next turn lands on a "
        "pod whose /workspace appears identical to the state at "
        "the end of the previous session.")

    add_h2(doc, "5.3", "Timing budget")

    add_para(doc,
        "On a warm path — existing claim, bound pod, restored "
        "workspace — the user-perceived time from message send "
        "to first streaming token is dominated by the Claude "
        "model latency (under 2 seconds at first-token, "
        "depending on prompt size). On a cold path — fresh "
        "claim, warm-pool pod bound but workspace empty — the "
        "addition is the restore-from-GCS step, which adds "
        "1-3 seconds depending on workspace size. On the coldest "
        "path — no warm-pool pod available — container start "
        "adds another 8-15 seconds. The warm pool size of 2 is "
        "the tuning knob that keeps the demo-tenant tail "
        "latency on the warm path.")


def section_6_deployment(doc):
    add_h1(doc, "6", "Deployment")

    add_h2(doc, "6.1", "GCP project layout")

    add_para(doc,
        "The system runs in a single GCP project. The reference "
        "deployment uses cpe-slarbi-nvd-ant-demos in us-central1. "
        "Services enabled: Cloud Run, GKE Autopilot, Artifact "
        "Registry, Firestore (native mode, named database "
        "cc-on-ge), Cloud Storage, Vertex AI (with the global "
        "endpoint configured for Claude), Cloud Build, Discovery "
        "Engine, and IAM Service Account Credentials API. No "
        "VPC service controls perimeter is enforced in v1, "
        "though the architecture supports adding one without "
        "structural change.")

    add_h2(doc, "6.2", "Terraform structure")

    add_table_styled(doc,
        headers=["File", "What it provisions"],
        rows=[
            ["main.tf",
             "Provider config, locals, top-level outputs."],
            ["gke.tf",
             "Autopilot cluster cc-sandbox with the Agent Sandbox "
             "addon enabled."],
            ["cloudrun.tf",
             "cc-a2a-bridge Cloud Run service, image tag, env, "
             "min-instances, ingress."],
            ["firestore.tf",
             "Named database cc-on-ge, region, indexes for "
             "sessions and memory."],
            ["storage.tf",
             "Workspace bucket with uniform IAM, lifecycle for "
             "old snapshots, retention settings."],
            ["iam.tf",
             "Service accounts (cc-a2a-bridge, cc-a2a-backend, "
             "cc-sandbox-sa, cc-a2a-builder), role bindings, "
             "Workload Identity bindings."],
            ["k8s.tf",
             "SandboxTemplate, SandboxWarmPool, NetworkPolicies, "
             "sandbox-router Service and Deployment."],
        ],
        col_widths_in=[1.6, 5.2],
    )

    add_h2(doc, "6.3", "Build pipeline")

    add_para(doc,
        "Two images are built — cc-a2a-bridge (the Cloud Run "
        "service) and cc-backend (the GKE pod). Both are built "
        "with Cloud Build via infra/cloudbuild.yaml, which is "
        "parameterized by _IMAGE and _TAG substitutions so the "
        "same yaml drives both targets. Images are pushed to "
        "us-central1-docker.pkg.dev/<project>/cc-on-ge/. Tag "
        "discipline is phase-numbered: phase8-r1, phase12-r5, "
        "and so on. Tags are immutable in practice — a new "
        "tag for a new image — so terraform applies that bump "
        "image tags produce identifiable revision-rollout "
        "events in Cloud Run and SandboxTemplate generation "
        "events on the GKE side.")

    add_h2(doc, "6.4", "Gated rollout pattern")

    add_para(doc,
        "Image bumps follow a five-step gate that has been "
        "refined through twelve phases. (1) Build the new image "
        "with Cloud Build, verify the build itself succeeded "
        "and the relevant smoke-test inside the Dockerfile "
        "passes. (2) Run a pip-freeze (or equivalent) inside "
        "the built image to verify dependency versions match "
        "expectations — the lesson from the May 2026 outage is "
        "that the image must be inspected directly, not just "
        "the requirements file. (3) Edit the terraform image "
        "tag and run terraform plan; verify the plan is exactly "
        "the expected number of in-place updates with no "
        "collateral drift. (4) Eyeball the plan and apply on "
        "explicit confirmation. (5) Watch logs for the first "
        "6 minutes after rollout, specifically waiting for "
        "periodic background jobs (the sweeper) to fire at "
        "least once successfully before declaring the deploy "
        "smoke-green. (6) Verify in a real Gemini Enterprise "
        "thread.")

    add_callout(doc, "LESSON",
        "The platform's \"Ready=True\" signal is upstream of the "
        "actual user-visible failure surface. Production-readiness "
        "requires a synthetic request that exercises the "
        "dependency-critical path, not just a probe response. "
        "See Section 7 for the principle.")

    add_h2(doc, "6.5", "Deploying into a fresh GCP project")

    add_para(doc,
        "The rough sequence to stand the system up in an empty "
        "project: enable the APIs listed in 6.1; create the "
        "Terraform state bucket and builder service account "
        "(scripts/bootstrap.sh); run the parameterized Cloud "
        "Build for both images to populate Artifact Registry; "
        "terraform apply to provision infrastructure; "
        "scripts/register-agent.sh to register with Discovery "
        "Engine; verify smoke and the three CUJs. The "
        "make bootstrap and make deploy targets package this "
        "for a reproducible setup, but the underlying steps are "
        "discoverable and the apply is gated rather than "
        "single-shot.")


def section_7_best_practices(doc):
    add_h1(doc, "7", "Best Practices and Deployment Learnings")

    add_para(doc,
        "These principles emerged from the twelve-phase build "
        "and are documented here as actionable guidance for "
        "anyone building a similar system. Each is named, "
        "explained, and grounded in the specific incident or "
        "decision that drove the lesson home.")

    # Principle 1
    add_h2(doc, "7.1", "Separation of orchestrator and executor (the "
                       "Option-A bet)")

    add_para(doc,
        "The orchestrator owns user identity, session state, "
        "memory, artifact emission, and the surface registered "
        "with the host platform. The executor (Claude Code, in "
        "this build) is a stateless subprocess invoked by one "
        "tool with a narrow contract. Mixing these — letting the "
        "executor become the entry point — produces a system "
        "where every architectural concern is coupled to the "
        "harness implementation. Adding a second harness later "
        "is a tool registration, not a rewrite. This principle "
        "is what makes \"add Antigravity\" a Phase 13 conversation "
        "rather than a redesign.")

    # Principle 2
    add_h2(doc, "7.2", "Real-path verification over synthetic green tests")

    add_para(doc,
        "Tests that fire a request through the same path the user "
        "fires reveal categories of failure that probe-only "
        "tests never see. The May 2026 outage in this system had "
        "every platform-level health check green — Cloud Run "
        "Ready=True, ContainerHealthy=True, /healthz returning "
        "200 every 30 seconds — while every real A2A request "
        "crashed inside the request handler with a K8s 401. The "
        "fix in deployment process was to add a synthetic "
        "request that exercises a SandboxClaim CRUD operation "
        "before declaring a deploy green. The general principle: "
        "container health is necessary but not sufficient.")

    # Principle 3
    add_h2(doc, "7.3", "Conservative defaults for unverified surfaces")

    add_para(doc,
        "When the behavior of an external surface is not "
        "deterministically known, default to the safer fallback. "
        "The Gemini Enterprise MIME allowlist is enumerated by "
        "trial — types not on the allowlist must be routed to "
        "Path B (signed URL) rather than attempting Path A and "
        "hoping it renders. Pattern: any time the system depends "
        "on a downstream surface that may evolve, choose the "
        "default that produces a working result on the broadest "
        "set of inputs, and treat the optimization (Path A) as a "
        "type-narrowed special case.")

    # Principle 4
    add_h2(doc, "7.4", "Audit the bug class, not just the reported instance")

    add_para(doc,
        "When a bug is identified — for example, the "
        "function-calling string-vs-int crash in "
        "read_workspace_file — the immediate fix in one tool is "
        "necessary but not sufficient. The pattern (function "
        "calling delivers typed arguments as strings) likely "
        "affects every other tool with non-string parameters. "
        "The May 2026 hotfix sweep audited all five workspace "
        "tools and found a second related defect "
        "(delete_workspace_file's confirm argument being "
        "string-truthy-evaluated to True for the string \"False\"), "
        "which would have caused silent hard-deletes. Audit "
        "siblings; do not fix only the one that paged.")

    # Principle 5
    add_h2(doc, "7.5", "Safety mechanisms need tests proving the unsafe path is closed")

    add_para(doc,
        "A safety guarantee that is enforced only by code review "
        "is no guarantee at all. The per-user storage isolation "
        "in this system is enforced by the IAM Credential Access "
        "Boundary; the test that proves it is make iso-test, "
        "which constructs a pod with user_key A's downscoped "
        "token, attempts to read users/B/, and asserts a 403 "
        "from storage.googleapis.com (not from application "
        "code). The test fails loudly if any future change "
        "accidentally grants the backend SA direct bucket access. "
        "If the safety property cannot be expressed as a test, "
        "the property is not in production.")

    # Principle 6
    add_h2(doc, "7.6", "Function-calling delivers typed args as strings; coerce at the boundary")

    add_para(doc,
        "JSON-schema function-calling describes the contract — "
        "this argument is an integer, this one is a boolean — "
        "but the wire format and the model's transcription "
        "frequently produce stringified values. Tools should "
        "coerce at the boundary, before validation logic runs. "
        "The pattern is small (int(arg) inside a try/except, "
        "explicit string-truthiness for booleans) but it is "
        "load-bearing. The May 2026 workspace-tools hotfix is the "
        "evidence of why.")

    # Principle 7
    add_h2(doc, "7.7", "Container health is necessary but not sufficient")

    add_para(doc,
        "(Lesson 8.20 in the project's internal lessons-learned "
        "tracking.) Cloud Run reporting Ready=True only means the "
        "platform launched the container and a /healthz probe "
        "responded. It says nothing about whether the binary can "
        "actually do its job. The May 2026 outage had every "
        "platform metric green while every real request crashed. "
        "Production-readiness requires a synthetic request that "
        "exercises the dependency-critical path before declaring "
        "deploy success.")

    # Principle 8
    add_h2(doc, "7.8", "Unpinned dependencies are a time bomb")

    add_para(doc,
        "(Lesson 8.19.) A reproducible build is not the same as a "
        "deterministic build. If a dependency list contains "
        "package names without == pins, two builds of identical "
        "source code can produce binaries with different "
        "transitive trees, and the difference can break production "
        "silently. The May 2026 incident here: bridge "
        "requirements.txt had kubernetes and google-auth "
        "unpinned, pip resolved kubernetes-36 (released that "
        "month, with a new aiohttp-based auth stack) when the "
        "prior known-good build had resolved kubernetes-35, and "
        "every SandboxClaim CRUD operation returned 401 from the "
        "GKE control plane. Pin precisely; document with dated "
        "rationale.")

    # Principle 9
    add_h2(doc, "7.9", "Pin SOURCE matters as much as pin VALUE")

    add_para(doc,
        "(Lesson 8.21.) A version pin in requirements.txt does "
        "nothing if the Dockerfile installs packages with bare "
        "names (RUN pip install \"kubernetes\"). In the recovery "
        "from the May 2026 outage, the first attempt pinned the "
        "versions in requirements.txt and rebuilt. The image "
        "still resolved to kubernetes-36 because the Dockerfile "
        "had per-package RUN lines that bypassed the requirements "
        "file entirely. The verification step that caught this — "
        "running pip freeze inside the built image and "
        "diffing against the working baseline — is the discipline "
        "the general principle teaches: trust the artifact, not "
        "the manifest.")

    # Principle 10
    add_h2(doc, "7.10", "Gated rollouts with eyeball discipline at every blast-radius boundary")

    add_para(doc,
        "Every change with non-local blast radius is gated. "
        "Image builds, terraform plans, and applies are separate "
        "explicit steps. Plans are eyeballed for unexpected drift "
        "before apply. Real-GE verification is performed by the "
        "engineer, not asserted by automation, before commit. "
        "The discipline costs minutes; the savings, when a "
        "subtle terraform-plan or build-output surprise is "
        "caught before apply, are hours of rollback work. The "
        "May 2026 outage's recovery, where the broken image was "
        "deployed because the smoke gate only checked container "
        "health rather than a real K8s call, is the case study.")

    # Principle 11
    add_h2(doc, "7.11", "Probe before integrate (the A2UI lesson)")

    add_para(doc,
        "When a new external surface is on the roadmap — A2UI v0.8 "
        "for native interactive UI in chat — the right first step "
        "is a minimum-viable probe that confirms the surface "
        "actually behaves as documented before any production "
        "wiring depends on it. The probe artifacts in scripts/"
        "phase12 were never merged into production paths; they "
        "stood up enough scaffolding to verify A2UI surfacing in "
        "GE, captured the result, and let the production path "
        "stay on the prior working version. Integration commits "
        "to the dependency; probes do not. The discipline "
        "prevents a category of \"I'll just use the new thing\" "
        "outages.")


def section_8_operational(doc):
    add_h1(doc, "8", "Operational Considerations")

    add_h2(doc, "8.1", "Observability")

    add_para(doc,
        "Per-turn cost tracking lives in the orchestrator's "
        "interpretation of the claude-agent-sdk ResultMessage. "
        "Each invocation of claude_code emits, at end of turn, a "
        "structured result that includes input_tokens, "
        "output_tokens, and the model used. The orchestrator "
        "writes those fields to the ADK event log; the bridge "
        "surfaces them in the A2A response metadata so GE's "
        "billing surface can attribute cost per turn and per user. "
        "Long-term observability is intentionally minimal in v1 — "
        "logs to Cloud Logging, no application metrics — and a "
        "Phase 14 tracing initiative is queued.")

    add_h2(doc, "8.2", "Logging surface — what is logged and what is never logged")

    add_para(doc,
        "Logs are deliberately scoped to operational signals. "
        "Every A2A request logs context_id, task_id, user_key, "
        "and elapsed time. Tool invocations log the tool name "
        "and a structured outcome (success / failure). Errors "
        "log full traceback. Two things are explicitly never "
        "logged: signed URLs (because the URL is sensitive and "
        "anyone with log access could otherwise retrieve files "
        "they should not see), and user PII (the user_key hash "
        "is the only identifier in the log stream; the original "
        "email or subject appears only inside the bridge's token "
        "resolution and is not persisted).")

    add_h2(doc, "8.3", "IAM hygiene patterns (grant-then-revoke)")

    add_para(doc,
        "Operational IAM grants — debugging access, smoke-script "
        "permissions, one-off troubleshooting — follow a "
        "grant-then-revoke pattern. The grant is scripted with an "
        "explicit revoke script alongside; the operator running "
        "the script knows up front what they will undo. The "
        "scripts/iam-hygiene.sh helpers, introduced in commit "
        "74e6cc9, codify this discipline. The general pattern: "
        "production IAM grants in the steady state are "
        "Terraform-managed; ad-hoc grants are ephemeral and have "
        "an explicit cleanup partner.")

    add_h2(doc, "8.4", "Warm-pool tuning")

    add_para(doc,
        "The pool size is the dominant tail-latency knob. Two "
        "pre-spawned pods is enough for the demo tenant; a "
        "production tenant should size to expected concurrent "
        "first-turn users. The pool can be scaled with "
        "kubectl edit sandboxwarmpool cc-backend-warm or via "
        "Terraform. Pool members consume node capacity even while "
        "idle, so the trade-off is cost-per-warm-pod versus "
        "first-turn latency. The idle-teardown threshold (30 "
        "minutes default) is a separate knob trading user-perceived "
        "session continuity for node-hour cost.")

    add_h2(doc, "8.5", "Failure modes and how they manifest")

    add_table_styled(doc,
        headers=["Failure", "Symptom in GE", "Where to look"],
        rows=[
            ["Bridge cannot reach K8s API",
             "Empty reply with \"?\" badge",
             "Cloud Logging cc-a2a-bridge service — look for 401 "
             "from kubernetes.client"],
            ["Pod fails to bind from warm pool",
             "Long delay then \"Couldn't connect\" error",
             "kubectl describe sandboxclaim cc-u-<key>; "
             "kubectl get pods -n cc-sandbox"],
            ["Workspace restore fails",
             "Pod responsive but workspace empty",
             "Pod logs for the restore step; "
             "gs://<bucket>/users/<key>/current/ for the manifest"],
            ["Path A artifact silently dropped",
             "Agent claims success, no chip appears",
             "MIME is not on the allowlist — should have been "
             "routed Path B"],
            ["Path B URL renders inline in browser",
             "Page opens instead of downloading",
             "Check signed-URL params for "
             "response-content-disposition=attachment"],
            ["Vertex inference fails",
             "Empty thinking pane, error event in stream",
             "VERTEXAI_PROJECT, CLOUD_ML_REGION env in pod; "
             "Vertex quota in the GCP project"],
        ],
        col_widths_in=[2.0, 2.3, 2.5],
    )

    add_h2(doc, "8.6", "Rollback discipline")

    add_para(doc,
        "Rollbacks are first-class operations. Terraform's "
        "image-tag-based revision model means rolling back a "
        "bridge or backend image is a symmetric edit-plan-apply "
        "of the same form as the forward bump. Cloud Run "
        "preserves prior revisions indefinitely (or until "
        "manually deleted), so the rollback target is always "
        "available. The May 2026 outage was resolved by exactly "
        "this pattern: revert two image tags in terraform, "
        "show the plan, apply, restoration verified within "
        "15 minutes of incident declaration. Rollback is not a "
        "punishment for deploying a bad image; it is the "
        "expected response while the fix-forward is being "
        "prepared.")


def section_9_roadmap(doc):
    add_h1(doc, "9", "Extensibility and Roadmap")

    add_h2(doc, "9.1", "Why Option A makes future additions additive")

    add_para(doc,
        "Each item on the roadmap below is a new tool or a new "
        "tightening, not a new architecture. This is the dividend "
        "of the orchestrator-as-brain separation. The cost of "
        "each addition is bounded by its own surface; none "
        "requires touching identity propagation, sandbox "
        "lifecycle, signed-URL minting, or the artifact-emission "
        "model.")

    add_h2(doc, "9.2", "Per-user authorization via authorizationConfig")

    add_para(doc,
        "Version 1 ships in single-user demo mode. The Gemini "
        "Enterprise Discovery Engine API supports a "
        "per-agent authorizationConfig surface that routes the "
        "user through a per-app OAuth consent screen and mints "
        "application-specific tokens for the agent. Integrating "
        "this is a bridge-side change: the bridge already derives "
        "user_key from any verified token; switching from the "
        "default GE token to an application-specific token "
        "changes nothing downstream. Phase 13 work item.")

    add_h2(doc, "9.3", "Latency optimization (model routing)")

    add_para(doc,
        "Today both the orchestrator and the Claude Code "
        "subprocess use claude-opus-4-7. For many user turns — "
        "particularly the recall-only follow-up questions and "
        "the workspace-listing requests — opus-class capability "
        "is excess. A model-routing tier in the orchestrator "
        "would inspect the turn shape and pick a smaller model "
        "(sonnet-4-6 or haiku-4-5) for the fast path. This is "
        "the original Phase 11 work item, currently queued. Net "
        "effect: an order-of-magnitude latency reduction on "
        "conversational turns without affecting code-generation "
        "quality.")

    add_h2(doc, "9.4", "Connector context ingestion")

    add_para(doc,
        "Gemini Enterprise ships native connectors for Drive, "
        "Docs, Jira, GitHub, and Confluence. Surfacing connector "
        "content as files under /workspace/context/ — read-only, "
        "scoped per user, refreshed on demand — would let users "
        "ask the agent to operate on documents they already "
        "have in those systems without uploading them. The "
        "implementation is a new tool, fetch_context, that the "
        "orchestrator can invoke at the start of a turn. Out of "
        "scope for v1; queued for the post-MVP roadmap.")

    add_h2(doc, "9.5", "A2UI for native interactive UI in chat")

    add_para(doc,
        "Gemini Enterprise's A2UI v0.8 protocol allows agents to "
        "emit structured UI components — buttons, forms, "
        "selectable choices — that render natively in the chat. "
        "The probe artifacts in scripts/phase12 (commit 5bb2e37) "
        "stood up enough scaffolding to verify the surface "
        "renders correctly. Production integration is deferred; "
        "the verified probe gives us confidence the surface is "
        "stable enough to integrate when the user value "
        "justifies it (likely the deploy-confirmation and "
        "schema-pick steps of artifact generation).")

    add_h2(doc, "9.6", "Live-preview proxy for dev servers")

    add_para(doc,
        "For prototype generation, the user often wants to see "
        "the prototype running, not just download the source. A "
        "live-preview proxy that exposes localhost:<port> from "
        "inside the sandbox over a signed-URL-authenticated "
        "tunnel would let the agent's working pane include an "
        "iframe of the running app. The architectural pieces "
        "exist (downscoped tokens, signed URLs); the integration "
        "is a new endpoint on the bridge and a tool on the "
        "orchestrator. Phase 15 candidate.")

    add_h2(doc, "9.7", "Observability and tracing")

    add_para(doc,
        "OpenTelemetry tracing across the bridge → router → pod → "
        "Claude Code chain would give us per-turn latency "
        "attribution at each hop. The trace context already "
        "propagates through A2A; instrumenting the bridge's "
        "FastAPI routes and the backend's ADK orchestrator with "
        "OTel spans, exporting to Cloud Trace, is a half-day "
        "of work. The benefit is the ability to answer \"why "
        "did that turn take 14 seconds\" with a flamegraph "
        "rather than log spelunking.")


def section_10_appendix(doc):
    add_h1(doc, "10", "Appendix")

    add_h2(doc, "A", "Locked configuration constants")

    add_table_styled(doc,
        headers=["Setting", "Value"],
        rows=[
            ["GCP project (reference)", "cpe-slarbi-nvd-ant-demos"],
            ["Region", "us-central1"],
            ["Model (orchestrator + tool)", "claude-opus-4-7"],
            ["Vertex endpoint", "global (CLOUD_ML_REGION=global)"],
            ["Claude Code on Vertex flag",
             "CLAUDE_CODE_USE_VERTEX=1"],
            ["Workspace bucket",
             "<project>-cc-a2a-snapshots"],
            ["Firestore database (named)", "cc-on-ge"],
            ["Sandbox claim name", "cc-u-<user_key>"],
            ["Sandbox port", "9000"],
            ["Idle teardown threshold", "30 min"],
            ["Cloud Run request timeout", "1 hour"],
            ["/workspace size", "20 Gi standard-rwo"],
            ["/workspace mount", "noexec"],
            ["Warm pool size", "2"],
            ["Node.js (for bundled Claude Code CLI)", "20"],
        ],
        col_widths_in=[3.0, 3.8],
    )

    add_h2(doc, "B", "Verified Gemini Enterprise MIME allowlist (Path A)")

    add_para(doc,
        "These MIME types render as native chips in Gemini "
        "Enterprise (verified by direct test). Everything else "
        "routes through Path B.")

    add_table_styled(doc,
        headers=["MIME", "Typical extension"],
        rows=[
            ["image/png", ".png"],
            ["image/jpeg", ".jpg / .jpeg"],
            ["application/pdf", ".pdf"],
            ["text/csv", ".csv"],
            ["text/plain", ".txt"],
            ["application/json", ".json"],
            ["application/vnd.openxmlformats-officedocument."
             "wordprocessingml.document",
             ".docx"],
            ["application/vnd.openxmlformats-officedocument."
             "spreadsheetml.sheet",
             ".xlsx"],
            ["application/vnd.openxmlformats-officedocument."
             "presentationml.presentation",
             ".pptx"],
        ],
        col_widths_in=[5.4, 1.4],
    )

    add_h2(doc, "C", "Dependency pin rationale")

    add_para(doc,
        "The bridge image pins two transitive dependencies "
        "explicitly on the Dockerfile RUN lines (the actual "
        "install source). The pins are recorded in "
        "requirements.txt as parallel documentation but are "
        "load-bearing on the Dockerfile.")

    add_table_styled(doc,
        headers=["Package", "Pin", "Rationale"],
        rows=[
            ["kubernetes", "==35.0.0",
             "kubernetes-36 (released May 2026) introduced an "
             "aiohttp-based async client whose auth wiring is "
             "rejected by the GKE control plane, producing 401 "
             "on every SandboxClaim CRUD. Re-test before "
             "bumping."],
            ["google-auth", "==2.52.0",
             "Precautionary, paired with the kubernetes pin "
             "because google-auth produces the token kubernetes "
             "uses. 2.52.0 was the version in the working "
             "phase8-r1 build."],
            ["a2a-sdk", "==0.2.13",
             "0.2.14 was yanked upstream (\"gRPC requirements "
             "handled incorrectly\"); 0.2.13 matches the "
             "a2a-protocol reference skill."],
        ],
        col_widths_in=[1.5, 1.1, 4.2],
    )

    add_h2(doc, "D", "CUJ acceptance criteria")

    add_para(doc,
        "A Critical User Journey is considered passing when all "
        "of the following are true for a real Gemini Enterprise "
        "thread test:")

    add_bullet(doc,
        "The user's prompt produces a streaming response visible "
        "in the working pane within 3 seconds of submission.")
    add_bullet(doc,
        "Tool calls render as collapsed cards in the working "
        "pane with their parameter values visible on expand.")
    add_bullet(doc,
        "The agent's text reply describes the artifact in plain "
        "language and includes either a native download chip "
        "(Path A) or a clickable link (Path B).")
    add_bullet(doc,
        "Clicking the chip or link downloads the file as an "
        "attachment — no inline browser rendering.")
    add_bullet(doc,
        "The downloaded file is well-formed for its declared "
        "MIME type (opens correctly in the default application).")
    add_bullet(doc,
        "A follow-up turn in the same thread successfully reads "
        "the file from workspace state and can iterate on it.")

    add_h2(doc, "E", "Glossary")

    add_table_styled(doc,
        headers=["Term", "Meaning"],
        rows=[
            ["A2A",
             "Agent-to-Agent protocol. The JSON-RPC over SSE "
             "protocol Gemini Enterprise Discovery Engine uses "
             "to call out to gallery agents."],
            ["ADK",
             "Agent Development Kit. Google's open-source Python "
             "framework for building tool-using agents."],
            ["Agent Sandbox",
             "GKE Autopilot addon (extensions.agents.x-k8s.io "
             "API group) that provides SandboxTemplate, "
             "SandboxClaim, and SandboxWarmPool CRDs for "
             "per-user pod lifecycle."],
            ["context_id",
             "A2A protocol identifier for a single GE chat "
             "thread. Maps 1:1 to an ADK Session in the backend "
             "pod."],
            ["CUJ",
             "Critical User Journey. The end-to-end user flow "
             "that defines acceptance for a feature or release."],
            ["Discovery Engine",
             "The Gemini Enterprise component that hosts the "
             "agent gallery and routes A2A requests."],
            ["Downscoped STS token",
             "A short-lived Google Cloud token derived from a "
             "service account credential and scoped to a "
             "narrower IAM surface via a Credential Access "
             "Boundary CEL expression."],
            ["FileWithBytes",
             "A2A response artifact containing base64-encoded "
             "file content. The Path A delivery mechanism."],
            ["Path A / Path B",
             "Project shorthand for the two artifact-delivery "
             "modes: inline FileWithBytes (Path A) vs signed "
             "Cloud Storage URL embedded in agent reply (Path B)."],
            ["SandboxClaim",
             "GKE Agent Sandbox CRD representing a user's "
             "binding to a sandbox pod. Created by the bridge as "
             "cc-u-<user_key>."],
            ["task_id",
             "A2A identifier for a single user turn within a "
             "context_id."],
            ["user_key",
             "Short SHA256 hash of the user's email or OAuth "
             "subject. The PII-free identifier used in K8s "
             "resource names, GCS prefixes, and Firestore "
             "documents."],
        ],
        col_widths_in=[1.8, 5.0],
    )


# ============================================================================
# Main orchestration
# ============================================================================

def main():
    doc = Document()

    # Section 1: cover (zero margins, no header/footer)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    build_cover(doc)

    # Page break to start a new section for the body
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_body_section(new_section)
    restart_page_numbers(new_section)

    # TOC
    build_toc(doc)
    add_page_break(doc)

    # Body sections 1-10
    section_1_executive(doc)
    add_page_break(doc)

    section_2_use_case(doc)
    add_page_break(doc)

    section_3_architecture(doc)
    add_page_break(doc)

    section_4_components(doc)
    add_page_break(doc)

    section_5_lifecycle(doc)
    add_page_break(doc)

    section_6_deployment(doc)
    add_page_break(doc)

    section_7_best_practices(doc)
    add_page_break(doc)

    section_8_operational(doc)
    add_page_break(doc)

    section_9_roadmap(doc)
    add_page_break(doc)

    section_10_appendix(doc)

    out = "/usr/local/google/home/slarbi/Documents/claude-code-agent/eddoc/cc-claude-code-on-ge-engineering-reference.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
